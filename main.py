
import os
import glob
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Create tmp folder
os.makedirs("tmp", exist_ok=True)

# App settings
st.set_page_config(page_title="AI-Powered Data Storyteller", layout="wide")
st.title("📊 AI-Powered Data Storyteller")

# ---------- Utility functions ----------

def auto_insights(df):
    insights = []
    nrows, ncols = df.shape
    insights.append(f"The dataset has {nrows} rows and {ncols} columns.")

    # Missing values
    miss = (df.isnull().mean() * 100).sort_values(ascending=False)
    for col, pct in miss.head(5).items():
        if pct > 0:
            insights.append(f"Column '{col}' has {pct:.0f}% missing values.")

    # Numeric overview
    num = df.select_dtypes(include="number")
    if not num.empty:
        top_var = num.var().sort_values(ascending=False).index[0]
        insights.append(f"Highest variance in numeric columns: {top_var}")
        corr = num.corr().abs().unstack().sort_values(ascending=False)
        corr = corr[corr < 1].drop_duplicates().head(3)
        for (a, b), v in corr.items():
            insights.append(f"Correlation {a} vs {b}: {v:.2f}")
    return insights

def generate_plots(df):
    # Clean up old plots
    for f in glob.glob("tmp/*.png"):
        os.remove(f)

    plot_paths = []

    # Bar chart for first categorical
    cat_cols = df.select_dtypes(exclude="number").columns
    if len(cat_cols) > 0:
        plt.figure(figsize=(6,4))
        df[cat_cols[0]].value_counts().head(10).plot.bar()
        plt.title(f"Top categories in {cat_cols[0]}")
        path = "tmp/bar.png"
        plt.savefig(path); plt.close()
        plot_paths.append(path)

    # Line chart for first numeric vs index
    num_cols = df.select_dtypes(include="number").columns
    if len(num_cols) > 0:
        plt.figure(figsize=(6,4))
        df[num_cols[0]].head(50).plot.line()
        plt.title(f"Trend of {num_cols[0]}")
        path = "tmp/line.png"
        plt.savefig(path); plt.close()
        plot_paths.append(path)

    # Heatmap of correlations
    if len(num_cols) > 1:
        corr = df[num_cols].corr()
        plt.figure(figsize=(6,4))
        sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm")
        plt.title("Correlation Heatmap")
        path = "tmp/heatmap.png"
        plt.savefig(path); plt.close()
        plot_paths.append(path)

    return plot_paths

def generate_report(insights, plot_paths):
    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    for i in insights:
        doc.add_paragraph(i)

    doc.add_heading("Key Charts", level=1)
    for p in plot_paths:
        doc.add_picture(p, width=Inches(5))

    # Save to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- Streamlit App Workflow ----------

uploaded_file = st.file_uploader("📁 Upload a CSV file", type="csv")

if uploaded_file:
    try:
        df = pd.read_csv(uploaded_file)
    except Exception as e:
        st.error(f"❌ Failed to read CSV file: {e}")
        st.stop()

    st.write("### 👀 Preview of Data")
    st.dataframe(df.head())

    if df.shape[0] < 2 or df.shape[1] < 1:
        st.error("❌ Dataset too small for analysis.")
    else:
        # Generate insights
        st.write("### 🔎 Insights")
        with st.spinner("Generating insights..."):
            insights = auto_insights(df)
        for i in insights:
            st.write("•", i)

        # Generate plots
        st.write("### 📈 Plots")
        with st.spinner("Generating plots..."):
            plot_paths = generate_plots(df)
        for p in plot_paths:
            st.image(p)

        # Generate and download report
        if st.button("📄 Generate Executive Report"):
            with st.spinner("Creating report..."):
                report_buffer = generate_report(insights, plot_paths)
            st.success("✅ Report ready!")
            st.download_button("⬇️ Download Report", report_buffer, file_name="executive_report.docx")

